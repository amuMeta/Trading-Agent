#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter
基于json_to_markdown.py的两步转换：JSON → Markdown → DOCX
"""

import os
import sys
from pathlib import Path
from typing import Optional, List
import argparse
from datetime import datetime
import re

# 导入markdown转换器
try:
    from .json_to_markdown import JSONToMarkdownConverter
except ImportError:
    try:
        from json_to_markdown import JSONToMarkdownConverter
    except ImportError:
        from src.export.json_to_markdown import JSONToMarkdownConverter

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("❌ 缺少依赖包，请安装：pip install python-docx")
    sys.exit(1)


class MarkdownToDocxConverter:
    """Markdown转DOCX转换器"""
    
    def __init__(self, dump_dir: str = "src/dump", key_agents_only: bool = False):
        """初始化转换器
        
        Args:
            dump_dir: dump文件夹路径
            key_agents_only: 是否只导出关键智能体
        """
        self.dump_dir = Path(dump_dir)
        self.key_agents_only = key_agents_only
        # 输出到 dumptools/docx_reports/ 目录
        self.output_dir = Path(__file__).parent / "docx_reports"
        self.output_dir.mkdir(exist_ok=True)
        
        # 初始化Markdown转换器
        self.md_converter = JSONToMarkdownConverter(str(self.dump_dir), key_agents_only=self.key_agents_only)
        
        # 目录项列表
        self.toc_entries = []
    
    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10)
        
        # 创建中文标题样式
        if 'Chinese Title' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('Chinese Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = '微软雅黑'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(20)
        
        # 创建各级标题样式
        heading_configs = [
            ('Chinese Heading 1', 16, True),
            ('Chinese Heading 2', 14, True),
            ('Chinese Heading 3', 12, True),
            ('Chinese Heading 4', 11, True)
        ]
        
        for style_name, font_size, is_bold in heading_configs:
            if style_name not in [s.name for s in doc.styles]:
                heading_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_font = heading_style.font
                heading_font.name = '微软雅黑'
                heading_font.size = Pt(font_size)
                heading_font.bold = is_bold
                heading_style.paragraph_format.space_after = Pt(10)
                heading_style.paragraph_format.space_before = Pt(10)
        
        # 创建代码样式
        if 'Chinese Code' not in [s.name for s in doc.styles]:
            code_style = doc.styles.add_style('Chinese Code', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
        
        # 创建引用样式
        if 'Chinese Quote' not in [s.name for s in doc.styles]:
            quote_style = doc.styles.add_style('Chinese Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_font = quote_style.font
            quote_font.name = '微软雅黑'
            quote_font.size = Pt(10)
            quote_font.italic = True
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.right_indent = Inches(0.5)
        
        # 创建目录样式
        if 'TOC Heading' not in [s.name for s in doc.styles]:
            toc_heading_style = doc.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_heading_font = toc_heading_style.font
            toc_heading_font.name = '微软雅黑'
            toc_heading_font.size = Pt(16)
            toc_heading_font.bold = True
            toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_heading_style.paragraph_format.space_after = Pt(20)
    
    def _add_emoji_support(self, doc):
        """添加emoji字体支持"""
        try:
            # 为文档添加Segoe UI Emoji字体支持
            # 这需要在运行时设置字体回退
            pass
        except Exception as e:
            print(f"⚠️ Emoji字体设置失败: {e}")
    
    def _create_table_of_contents(self, doc):
        """创建目录"""
        if not self.toc_entries:
            return
        
        # 添加目录标题
        toc_heading = doc.add_paragraph('📋 目录', style='TOC Heading')
        
        # 添加目录项
        for level, title, page_num in self.toc_entries:
            toc_paragraph = doc.add_paragraph()
            
            # 设置缩进
            toc_paragraph.paragraph_format.left_indent = Inches((level - 1) * 0.3)
            
            # 添加标题文本
            run = toc_paragraph.add_run(title)
            run.font.name = '微软雅黑'
            
            # 尝试设置emoji字体
            if any(ord(char) > 127 for char in title if ord(char) > 0x1F000):
                run.font.name = 'Segoe UI Emoji'
            
            # 添加点线和页码（简化版）
            dots_run = toc_paragraph.add_run(' ' + '.' * (50 - len(title)))
            dots_run.font.name = '微软雅黑'
            
            page_run = toc_paragraph.add_run(f' {page_num}')
            page_run.font.name = '微软雅黑'
        
        # 添加分页符
        doc.add_page_break()
    
    def _parse_markdown_to_docx(self, markdown_content: str, doc):
        """解析Markdown内容并添加到DOCX文档"""
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading = line.lstrip('#').strip()
                
                if level == 1:
                    p = doc.add_paragraph(heading, style='Chinese Heading 1')
                elif level == 2:
                    p = doc.add_paragraph(heading, style='Chinese Heading 2')
                elif level == 3:
                    p = doc.add_paragraph(heading, style='Chinese Heading 3')
                else:
                    p = doc.add_paragraph(heading, style='Chinese Heading 4')
                
                # 为emoji设置特殊字体
                for run in p.runs:
                    if any(ord(char) > 0x1F000 for char in run.text):
                        run.font.name = 'Segoe UI Emoji'
            
            # 处理引用
            elif line.startswith('> '):
                quote = line[2:].strip()
                p = doc.add_paragraph(quote, style='Chinese Quote')
            
            # 处理列表
            elif line.startswith('- '):
                item = line[2:].strip()
                p = doc.add_paragraph(item, style='List Bullet')
                # 设置字体
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
            
            # 处理表格
            elif '|' in line and line.strip().startswith('|'):
                # 收集表格行
                table_rows = []
                table_rows.append(line.strip())
                
                # 继续读取表格行
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if next_line and '|' in next_line and next_line.startswith('|'):
                        table_rows.append(next_line)
                        j += 1
                    else:
                        break
                
                # 解析表格并添加到文档
                if len(table_rows) > 1:
                    self._add_table_to_doc(table_rows, doc)
                    i = j - 1  # 调整索引
            
            # 处理代码块
            elif line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    # 保持完整内容，不进行截断
                    
                    # 分行添加代码
                    for code_line in code_text.split('\n'):
                        p = doc.add_paragraph(code_line, style='Chinese Code')
            
            # 处理普通文本
            else:
                if line:
                    p = doc.add_paragraph()
                    
                    # 处理粗体文本
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            # 粗体文本
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            # 普通文本
                            run = p.add_run(part)
                        
                        # 设置字体
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
            
            i += 1
    
    def _add_table_to_doc(self, table_rows, doc):
        """将表格添加到Word文档"""
        # 解析表格数据
        table_data = []
        for row in table_rows:
            # 移除首尾的|符号并分割
            cells = [cell.strip() for cell in row.strip('|').split('|')]
            table_data.append(cells)
        
        # 跳过分隔行（通常是第二行，包含---）
        if len(table_data) > 1 and all('---' in cell or '-' in cell for cell in table_data[1]):
            table_data.pop(1)
        
        if table_data and len(table_data) > 0:
            # 创建表格
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            
            if cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # 填充表格数据
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            
                            # 设置字体
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'
                                    run.font.size = Pt(9)
                            
                            # 表头加粗
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                # 添加段落间距
                doc.add_paragraph()

    def convert_json_to_docx_via_markdown(self, json_file_path: str) -> Optional[str]:
        """通过Markdown中间步骤将JSON转换为DOCX
        
        Args:
            json_file_path: JSON文件路径
            
        Returns:
            生成的DOCX文件路径，失败返回None
        """
        try:
            # 第一步：JSON转Markdown
            print(f"📄 第一步：将JSON转换为Markdown...")
            md_file_path = self.md_converter.convert_json_to_markdown(json_file_path)
            
            if not md_file_path or not os.path.exists(md_file_path):
                print("❌ Markdown转换失败")
                return None
            
            # 第二步：Markdown转DOCX
            print(f"📄 第二步：将Markdown转换为DOCX...")
            
            # 读取Markdown文件
            with open(md_file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # 创建新的Word文档
            doc = Document()
            
            # 设置文档样式
            self._setup_document_styles(doc)
            self._add_emoji_support(doc)
            
            # 解析Markdown内容并添加到文档
            self._parse_markdown_to_docx(markdown_content, doc)
            
            # 生成DOCX文件名
            json_filename = Path(json_file_path).stem
            if self.key_agents_only:
                docx_file = self.output_dir / f"{json_filename}_关键分析.docx"
            else:
                docx_file = self.output_dir / f"{json_filename}.docx"
            
            # 保存文档
            doc.save(str(docx_file))
            
            print(f"✅ DOCX报告已生成: {docx_file}")
            return str(docx_file)
            
        except Exception as e:
            print(f"❌ 转换失败: {e}")
            return None
    
    def convert_latest_json(self) -> Optional[str]:
        """转换最新的JSON文件
        
        Returns:
            生成的DOCX文件路径，失败返回None
        """
        try:
            # 查找dump目录下的所有JSON文件
            json_files = list(self.dump_dir.glob("session_*.json"))
            
            if not json_files:
                print(f"❌ 在 {self.dump_dir} 目录下未找到JSON文件")
                return None
            
            # 找到最新的文件
            latest_json = max(json_files, key=lambda f: f.stat().st_mtime)
            print(f"📄 找到最新的JSON文件: {latest_json.name}")
            
            # 转换为DOCX
            return self.convert_json_to_docx_via_markdown(str(latest_json))
            
        except Exception as e:
            print(f"❌ 转换过程中发生错误: {e}")
            return None
    
    def convert_all_json(self) -> List[str]:
        """转换所有JSON文件
        
        Returns:
            生成的DOCX文件路径列表
        """
        try:
            # 查找dump目录下的所有JSON文件
            json_files = list(self.dump_dir.glob("session_*.json"))
            
            if not json_files:
                print(f"❌ 在 {self.dump_dir} 目录下未找到JSON文件")
                return []
            
            results = []
            for json_file in json_files:
                print(f"📄 转换文件: {json_file.name}")
                result = self.convert_json_to_docx_via_markdown(str(json_file))
                if result:
                    results.append(result)
            
            return results
            
        except Exception as e:
            print(f"❌ 批量转换过程中发生错误: {e}")
            return []


def main():
    """主函数 - 命令行工具"""
    parser = argparse.ArgumentParser(
        description="Markdown to DOCX Converter - 通过Markdown中间步骤将JSON转换为DOCX"
    )
    parser.add_argument("-f", "--file", help="指定要转换的JSON文件路径")
    parser.add_argument("-l", "--latest", action="store_true", help="转换最新的JSON文件")
    parser.add_argument("-a", "--all", action="store_true", help="转换所有JSON文件")
    parser.add_argument("-d", "--dump-dir", default="src/dump", help="dump文件夹路径")
    
    args = parser.parse_args()
    
    converter = MarkdownToDocxConverter(args.dump_dir)
    
    if args.all:
        # 转换所有文件
        results = converter.convert_all_json()
        if results:
            print(f"🎉 批量转换完成，共生成 {len(results)} 个DOCX文件")
        else:
            print("❌ 批量转换失败")
    
    elif args.latest:
        # 转换最新文件
        result = converter.convert_latest_json()
        if result:
            print(f"🎉 转换完成: {result}")
    
    elif args.file:
        # 转换指定文件
        if os.path.exists(args.file):
            result = converter.convert_json_to_docx_via_markdown(args.file)
            if result:
                print(f"🎉 转换完成: {result}")
        else:
            print(f"❌ 文件不存在: {args.file}")
    
    else:
        # 默认转换最新文件
        result = converter.convert_latest_json()
        if result:
            print(f"🎉 转换完成: {result}")


if __name__ == "__main__":
    main()